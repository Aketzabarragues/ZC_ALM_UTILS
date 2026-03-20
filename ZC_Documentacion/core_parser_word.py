import os
import re
import mammoth
import unicodedata
import tempfile
import core_logger as log

# ==============================================================================
# MAGIA DE COLORES: Método de Marcadores (Engañando a Mammoth)
# ==============================================================================
def preprocesar_colores_word(ruta_word):
    """
    Lee el Word original, inyecta marcas de color [COLOR:#RGB] en los textos
    y guarda un documento temporal. Así burlamos el filtro de Mammoth.
    """
    try:
        import docx
    except ImportError:
        log.error("Falta la librería python-docx para colores. Ejecuta: pip install python-docx")
        return ruta_word
        
    log.info("Analizando y protegiendo colores del documento Word...")
    doc = docx.Document(ruta_word)
    
    def procesar_runs(runs):
        for r in runs:
            # Si el fragmento de texto tiene color explícito (no automático/negro)
            if r.font.color and r.font.color.rgb:
                color_hex = str(r.font.color.rgb)
                # Ignoramos si es negro puro (000000) o no está definido
                if color_hex and color_hex != "000000" and color_hex != "None":
                    # Rodear el texto con nuestros marcadores seguros
                    if r.text.strip(): 
                        r.text = f"[COLOR:#{color_hex}]{r.text}[/COLOR]"

    # 1. Buscamos colores en párrafos normales
    for p in doc.paragraphs:
        procesar_runs(p.runs)
                
    # 2. Buscamos colores ocultos dentro de las tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar_runs(p.runs)
                    
    # Guardamos el documento modificado en los temporales de Windows
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, "temp_zcalm_color.docx")
    doc.save(temp_path)
    
    return temp_path
# ==============================================================================


def procesar_h3_en_contenido(html_contenido):
    """Detecta los H3, crea sus identificadores únicos y prepara las subsecciones."""
    subsecciones = []
    partes_h3 = re.split(r'(<h3\b[^>]*>.*?<\/h3>)', html_contenido, flags=re.IGNORECASE | re.DOTALL)
    nuevo_html = ""
    
    if partes_h3[0].strip():
        nuevo_html += f'<div class="doc-section">\n{partes_h3[0]}\n</div>\n'
        
    for i in range(1, len(partes_h3), 2):
        h3_full = partes_h3[i]
        contenido_h3 = partes_h3[i+1] if i+1 < len(partes_h3) else ""
        
        texto_h3 = re.sub('<[^<]+?>', '', h3_full).strip()
        texto_norm = unicodedata.normalize('NFKD', texto_h3).encode('ASCII', 'ignore').decode('utf-8')
        id_h3 = "sec_" + re.sub(r'[^a-zA-Z0-9]+', '_', texto_norm).strip('_')[:30]
        
        subsecciones.append({"id": id_h3, "titulo": texto_h3, "nivel": 3})
        nuevo_html += f'<div class="doc-section" id="{id_h3}">\n{h3_full}\n{contenido_h3}\n</div>\n'
        
    return nuevo_html, subsecciones

def trocear_html_por_niveles(html_puro):
    """Trocea el documento en capítulos (H1/H2)."""
    partes = re.split(r'(<(h[12])\b[^>]*>.*?<\/\2>)', html_puro, flags=re.IGNORECASE | re.DOTALL)
    capitulos = []
    
    if partes[0].strip():
        html_mod, subsecciones = procesar_h3_en_contenido(partes[0])
        capitulos.append({"archivo": "000_portada.html", "titulo": "Portada y Revisiones", "nivel": 1, "contenido": html_mod, "subsecciones": subsecciones})
    
    contador = 1
    for i in range(1, len(partes), 3):
        etiqueta_completa = partes[i]
        tipo_header = partes[i+1].lower() 
        contenido_subsecuente = partes[i+2] if i+2 < len(partes) else ""
        
        texto_limpio = re.sub('<[^<]+?>', '', etiqueta_completa).strip()
        texto_normalizado = unicodedata.normalize('NFKD', texto_limpio).encode('ASCII', 'ignore').decode('utf-8')
        nombre_seguro = re.sub(r'[^a-zA-Z0-9]+', '_', texto_normalizado).strip('_')[:50]
        
        contenido_total = etiqueta_completa + contenido_subsecuente
        html_mod, subsecciones = procesar_h3_en_contenido(contenido_total)
        
        capitulos.append({
            "archivo": f"{contador:03d}_{nombre_seguro}.html",
            "titulo": texto_limpio,
            "nivel": int(tipo_header[1]),
            "contenido": html_mod,
            "subsecciones": subsecciones
        })
        contador += 1
    return capitulos

def procesar_word(ruta_word, ruta_destino):
    """Extrae el Word, guarda imágenes, protege colores y devuelve capítulos."""
    log.info("Iniciando extracción del documento Word...")
    ruta_img = os.path.join(ruta_destino, 'img')
    contador_img = 1

    # 1. PRE-PROCESAMOS EL WORD (Inyectamos nuestras etiquetas seguras)
    ruta_temp_word = preprocesar_colores_word(ruta_word)

    def convertir_imagen(image):
        nonlocal contador_img
        extension = image.content_type.split("/")[1]
        nombre_img = f"img_{contador_img}.{extension}"
        with image.open() as image_stream:
            with open(os.path.join(ruta_img, nombre_img), "wb") as f:
                f.write(image_stream.read())
        contador_img += 1
        return {"src": f"../img/{nombre_img}"}

    # 2. MANDAMOS EL WORD TEMPORAL A MAMMOTH
    with open(ruta_temp_word, "rb") as docx_file:
        resultado = mammoth.convert_to_html(
            docx_file, 
            convert_image=mammoth.images.img_element(convertir_imagen)
        )
    
    html_crudo = resultado.value

    # 3. CONVERTIMOS NUESTROS MARCADORES A HTML REAL (SPAN COLOR)
    # Reemplaza [COLOR:#FF0000]texto[/COLOR] por <span style="color: #FF0000;">texto</span>
    html_con_colores = re.sub(
        r'\[COLOR:#([0-9a-fA-F]{6})\](.*?)\[/COLOR\]', 
        r'<span style="color: #\1;">\2</span>', 
        html_crudo, 
        flags=re.DOTALL
    )

    # Borramos el archivo temporal para no ensuciar el ordenador
    if ruta_temp_word != ruta_word and os.path.exists(ruta_temp_word):
        os.remove(ruta_temp_word)
    
    # 4. TROCEAMOS Y DEVOLVEMOS EL CONTENIDO YA COLOREADO
    capitulos = trocear_html_por_niveles(html_con_colores)
    
    return capitulos