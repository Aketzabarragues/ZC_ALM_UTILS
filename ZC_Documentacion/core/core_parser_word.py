"""
Analizador (Parser) y Conversor de Documentos Word.

Este módulo se encarga de procesar el manual corporativo en formato .docx,
transformándolo en una colección de vistas HTML estáticas.

Implementa un flujo híbrido que utiliza 'python-docx' para la extracción y 
preservación de metadatos de estilo (como colores RGB) que normalmente son 
descartados por motores semánticos puros, y 'mammoth' para la conversión 
robusta del documento a HTML estándar.
"""

import os
import re
import mammoth
import unicodedata
import tempfile
from core import core_logger as log


def preprocesar_colores_word(ruta_word):
    """
    Preprocesa el documento Word inyectando marcadores de retención de color.
    
    Dado que la librería Mammoth elimina intencionadamente los estilos manuales 
    para priorizar la semántica, esta función escanea el documento previamente 
    con python-docx. Al detectar texto con color explícito, lo envuelve en una 
    etiqueta segura (ej. [COLOR:#FF0000]texto[/COLOR]). Estas etiquetas 
    sobreviven a la conversión a texto plano de Mammoth y pueden ser 
    reemplazadas posteriormente por código CSS/HTML real.

    Args:
        ruta_word (str): Ruta absoluta al documento Word original.

    Returns:
        str: Ruta absoluta al documento Word temporal modificado.
    """
    try:
        import docx
    except ImportError:
        log.error("Dependencia faltante: python-docx requerida para retención de colores.")
        return ruta_word
        
    log.info("Analizando y protegiendo paleta de colores del documento Word...")
    doc = docx.Document(ruta_word)
    
    def procesar_runs(runs):
        """Inspecciona y modifica fragmentos de texto (runs) conservando su color."""
        for r in runs:
            if r.font.color and r.font.color.rgb:
                color_hex = str(r.font.color.rgb)
                # Exclusión de color automático (negro) o indefinido
                if color_hex and color_hex != "000000" and color_hex != "None":
                    if r.text.strip(): 
                        r.text = f"[COLOR:#{color_hex}]{r.text}[/COLOR]"

    # 1. Inspección de nodos de texto estándar (párrafos)
    for p in doc.paragraphs:
        procesar_runs(p.runs)
                
    # 2. Inspección recursiva de nodos de texto anidados (tablas)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar_runs(p.runs)
                    
    # Serialización del documento modificado en el directorio temporal del sistema OS
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, "temp_zcalm_color.docx")
    doc.save(temp_path)
    
    return temp_path


def procesar_h3_en_contenido(html_contenido):
    """
    Detecta y procesa subsecciones (etiquetas <h3>) dentro de un bloque HTML.
    
    Aplica identificadores únicos generados dinámicamente basados en el título
    normalizado, permitiendo la creación de anclas para enlaces profundos (Deep Linking).

    Args:
        html_contenido (str): Fragmento HTML a procesar.

    Returns:
        tuple: (HTML inyectado con clases e IDs, Lista de DTOs de subsecciones).
    """
    subsecciones = []
    # Segmentación inclusiva reteniendo la etiqueta delimitadora
    partes_h3 = re.split(r'(<h3\b[^>]*>.*?<\/h3>)', html_contenido, flags=re.IGNORECASE | re.DOTALL)
    nuevo_html = ""
    
    if partes_h3[0].strip():
        nuevo_html += f'<div class="doc-section">\n{partes_h3[0]}\n</div>\n'
        
    for i in range(1, len(partes_h3), 2):
        h3_full = partes_h3[i]
        contenido_h3 = partes_h3[i+1] if i+1 < len(partes_h3) else ""
        
        # Normalización ASCII para generación de identificadores URI-safe
        texto_h3 = re.sub('<[^<]+?>', '', h3_full).strip()
        texto_norm = unicodedata.normalize('NFKD', texto_h3).encode('ASCII', 'ignore').decode('utf-8')
        id_h3 = "sec_" + re.sub(r'[^a-zA-Z0-9]+', '_', texto_norm).strip('_')[:30]
        
        subsecciones.append({"id": id_h3, "titulo": texto_h3, "nivel": 3})
        nuevo_html += f'<div class="doc-section" id="{id_h3}">\n{h3_full}\n{contenido_h3}\n</div>\n'
        
    return nuevo_html, subsecciones


def trocear_html_por_niveles(html_puro):
    """
    Fragmenta el documento HTML monolítico en capítulos discretos.
    
    Utiliza expresiones regulares para dividir el flujo basado en cabeceras 
    de primer y segundo nivel (H1, H2), generando un DTO estructural por página.

    Args:
        html_puro (str): Documento HTML continuo exportado por Mammoth.

    Returns:
        list: Colección de diccionarios que representan los capítulos separados.
    """
    partes = re.split(r'(<(h[12])\b[^>]*>.*?<\/\2>)', html_puro, flags=re.IGNORECASE | re.DOTALL)
    capitulos = []
    
    if partes[0].strip():
        html_mod, subsecciones = procesar_h3_en_contenido(partes[0])
        capitulos.append({
            "archivo": "000_portada.html", 
            "titulo": "Portada y Revisiones", 
            "nivel": 1, 
            "contenido": html_mod, 
            "subsecciones": subsecciones
        })
    
    contador = 1
    for i in range(1, len(partes), 3):
        etiqueta_completa = partes[i]
        tipo_header = partes[i+1].lower() 
        contenido_subsecuente = partes[i+2] if i+2 < len(partes) else ""
        
        # Generación de nombre de archivo estandarizado seguro
        texto_limpio = re.sub('<[^<]+?>', '', etiqueta_completa).strip()
        texto_normalizado = unicodedata.normalize('NFKD', texto_limpio).encode('ASCII', 'ignore').decode('utf-8')
        nombre_seguro = re.sub(r'[^a-zA-Z0-9]+', '_', texto_normalizado).strip('_')[:50]
        
        contenido_total = etiqueta_completa + contenido_subsecuente
        html_mod, subsecciones = procesar_h3_en_contenido(contenido_total)
        
        capitulos.append({
            "archivo": f"{contador:03d}_{nombre_seguro}.html",
            "titulo": texto_limpio,
            "nivel": int(tipo_header[1]),
            "contenido": html_mod,
            "subsecciones": subsecciones
        })
        contador += 1
        
    return capitulos


def procesar_word(ruta_word, ruta_destino):
    """
    Controlador principal de extracción y conversión del manual Word.
    
    Fases del proceso:
    1. Inyección de metadatos de color (Preprocesamiento).
    2. Extracción de imágenes y conversión estructural a HTML semántico.
    3. Restauración de metadatos de color mediante CSS in-line.
    4. Fragmentación en páginas HTML discretas para el visor web.

    Args:
        ruta_word (str): Path de origen del documento .docx.
        ruta_destino (str): Path base donde se desplegarán las imágenes extraídas.

    Returns:
        list: Inventario completo de capítulos estructurados listos para enlazado.
    """
    log.info("Inicializando motor de extracción del documento Word...")
    ruta_img = os.path.join(ruta_destino, 'img')
    contador_img = 1

    # 1. Preprocesamiento: Inyección de anclas de retención para colores RGB
    ruta_temp_word = preprocesar_colores_word(ruta_word)

    def convertir_imagen(image):
        """Callback inyectado en Mammoth para la persistencia física de imágenes incrustadas."""
        nonlocal contador_img
        extension = image.content_type.split("/")[1]
        nombre_img = f"img_{contador_img}.{extension}"
        with image.open() as image_stream:
            with open(os.path.join(ruta_img, nombre_img), "wb") as f:
                f.write(image_stream.read())
        contador_img += 1
        return {"src": f"../img/{nombre_img}"}

    # 2. Conversión estructural delegando el volcado de imágenes al callback
    with open(ruta_temp_word, "rb") as docx_file:
        resultado = mammoth.convert_to_html(
            docx_file, 
            convert_image=mammoth.images.img_element(convertir_imagen)
        )
    
    html_crudo = resultado.value

    # 3. Postprocesamiento: Resolución de marcadores y restitución del formato CSS
    html_con_colores = re.sub(
        r'\[COLOR:#([0-9a-fA-F]{6})\](.*?)\[/COLOR\]', 
        r'<span style="color: #\1;">\2</span>', 
        html_crudo, 
        flags=re.DOTALL
    )

    # Higiene del sistema: Liberación del documento transaccional temporal
    if ruta_temp_word != ruta_word and os.path.exists(ruta_temp_word):
        os.remove(ruta_temp_word)
    
    # 4. Fragmentación algorítmica del documento continuo en páginas navegables
    capitulos = trocear_html_por_niveles(html_con_colores)
    
    return capitulos